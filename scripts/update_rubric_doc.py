"""Update GUIDE_Rubric_Document.docx from v3.0.0 to v3.1.2.

Adds the v3.0.2 ARCS work, v3.0.3 Three Alignments framing, and v3.1.0 Archetype 10
to the static rubric document without rebuilding it from scratch. Preserves the
existing Heading 1/2/3, Normal, List Paragraph styles and table formatting by
cloning existing elements and editing text rather than constructing new ones.

Run from the repo root:
    python scripts/update_rubric_doc.py
"""

from __future__ import annotations

import copy
import re
from pathlib import Path

import docx
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from lxml import etree

DOCX_PATH = Path("GUIDE_Rubric_Document.docx")


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_para_text(p, text: str) -> None:
    """Replace a paragraph's text while preserving its first run's style.

    python-docx paragraph.text setter wipes runs and formatting; this preserves
    the first run's font/style by mutating in place.
    """
    if not p.runs:
        p.text = text
        return
    p.runs[0].text = text
    # Remove any trailing runs
    for run in p.runs[1:]:
        run._element.getparent().remove(run._element)


def clone_para(template_p, new_text: str | None = None):
    """Clone a paragraph element. Optionally replace its text."""
    new_el = copy.deepcopy(template_p._element)
    new_p = docx.text.paragraph.Paragraph(new_el, template_p._parent)
    if new_text is not None:
        set_para_text(new_p, new_text)
    return new_p


def insert_para_after(target_p, new_p):
    """Insert new_p's XML element immediately after target_p in the body."""
    target_p._element.addnext(new_p._element)


def insert_para_before(target_p, new_p):
    target_p._element.addprevious(new_p._element)


def insert_table_after(target_p_or_t, new_t):
    """Insert table XML element after a paragraph or table element."""
    target_p_or_t._element.addnext(new_t._element)


def clone_table(template_table):
    """Deep-clone a table element. Returns python-docx Table wrapper."""
    new_el = copy.deepcopy(template_table._element)
    new_t = docx.table.Table(new_el, template_table._parent)
    return new_t


def find_para(doc, predicate):
    """Find the first paragraph matching predicate (style_name, text) -> bool."""
    for p in doc.paragraphs:
        if predicate(p.style.name, p.text):
            return p
    return None


def find_table_after_para(doc, para, max_skip=5):
    """Walk forward from para's element to find the next table sibling.

    Returns the python-docx Table wrapper or None.
    """
    el = para._element
    for _ in range(max_skip * 5):
        el = el.getnext()
        if el is None:
            return None
        tag = el.tag.split("}")[-1]
        if tag == "tbl":
            # Find the matching Table wrapper
            for t in doc.tables:
                if t._element is el:
                    return t
            return None
    return None


# ---------------------------------------------------------------------------
# Edit functions
# ---------------------------------------------------------------------------

def update_cover_version(doc):
    """Cover page: bump 'Version: 3.0.0 | April 2026' to 'Version: 3.1.2 | May 2026'."""
    for p in doc.paragraphs[:10]:
        if "3.0.0" in p.text and "April 2026" in p.text:
            set_para_text(p, "Version: 3.1.2 | May 2026")
            print(f"  cover version -> 3.1.2 | May 2026")
            return
    print("  WARN: cover version paragraph not found")


def update_executive_summary(doc):
    """Update 'nine evaluation archetypes' -> 'ten...' and '54 dimensions' references.

    Targets the Executive Summary paragraphs near the top.
    """
    edits = [
        ("nine evaluation archetypes", "ten evaluation archetypes"),
        ("nine archetypes and 54 dimensions", "ten archetypes and 60 dimensions"),
        ("nine archetypes are not arbitrary", "ten archetypes are not arbitrary"),
        ("Together, the nine archetypes and 54 dimensions",
         "Together, the ten archetypes and 60 dimensions"),
    ]
    for p in doc.paragraphs[:30]:
        for old, new in edits:
            if old in p.text:
                new_text = p.text.replace(old, new)
                set_para_text(p, new_text)
                print(f"  exec summary: '{old}' -> '{new}'")


def update_why_n_archetypes(doc):
    """Rename H2 'Why Nine Archetypes?' to 'Why Ten Archetypes?' and add a bullet
    for archetype 10 between the existing bullets and the ADDIE closing sentence."""
    # Find H2 heading
    h2 = find_para(doc, lambda style, text: style == "Heading 2" and "Why Nine" in text)
    if h2 is None:
        print("  WARN: 'Why Nine Archetypes?' heading not found")
        return
    set_para_text(h2, "Why Ten Archetypes?")
    print("  H2 renamed -> 'Why Ten Archetypes?'")

    # Find the existing bullets (List Paragraph style) under this section
    # and the closing 'Together,' paragraph
    bullets = []
    closing = None
    el = h2._element
    while True:
        el = el.getnext()
        if el is None:
            break
        tag = el.tag.split("}")[-1]
        if tag != "p":
            continue
        # Find matching python-docx paragraph
        p_match = next((p for p in doc.paragraphs if p._element is el), None)
        if p_match is None:
            continue
        if p_match.style.name == "Heading 2" or p_match.style.name == "Heading 1":
            break
        if p_match.style.name == "List Paragraph":
            bullets.append(p_match)
        elif p_match.text.startswith("Together,"):
            closing = p_match
            break

    if not bullets:
        print("  WARN: bullets under 'Why Nine' not found")
        return

    # Clone the last bullet to preserve List Paragraph styling, set new text,
    # and insert before the closing paragraph.
    new_bullet = clone_para(
        bullets[-1],
        "Archetype 10 (added in v3.1.0) is the synthesis lens. It evaluates whether "
        "the components from the other archetypes - objectives, instructional "
        "strategies, and assessments - cohere with each other and with broader "
        "curriculum and discipline scopes. It does not re-evaluate component "
        "quality; it asks whether the components, taken together, say the same thing.",
    )
    if closing is not None:
        insert_para_before(closing, new_bullet)
    else:
        insert_para_after(bullets[-1], new_bullet)
    print("  added archetype-10 bullet under 'Why Ten Archetypes?'")

    # Update the closing 'Together,' paragraph if found
    if closing is not None and "nine archetypes" in closing.text:
        set_para_text(closing, closing.text.replace("nine archetypes", "ten archetypes")
                                          .replace("54 dimensions", "60 dimensions"))
        print("  'Together,' closing paragraph updated")


def add_quality_outcomes_section(doc):
    """Insert a new H2 'Quality Outcomes: Effective, Efficient, Engaging' section
    after the 'Why Ten Archetypes?' section and before 'Quick-Reference'."""
    # Anchor: H2 'Quick-Reference: Archetypes at a Glance'
    anchor = find_para(doc, lambda style, text: style == "Heading 2" and "Quick-Reference" in text)
    if anchor is None:
        print("  WARN: 'Quick-Reference' anchor not found, cannot add Quality Outcomes")
        return

    # Need an H2 paragraph and several Normal paragraphs. Clone existing ones.
    existing_h2 = anchor  # Use 'Quick-Reference' H2 as the H2 template
    # For Normal paragraphs, use one near the H2 anchor
    existing_normal = next(
        (p for p in doc.paragraphs if p.style.name == "Normal" and p.text.strip()),
        None,
    )
    if existing_normal is None:
        print("  WARN: no Normal paragraph template found")
        return

    elements_to_insert = []

    # H2 heading
    h2 = clone_para(existing_h2, "Quality Outcomes: Effective, Efficient, Engaging")
    elements_to_insert.append(h2)

    # Intro paragraph
    p1 = clone_para(
        existing_normal,
        "High-quality instruction is effective, efficient, and engaging (Merrill, 2002). "
        "Hirumi (2025), building on Hirumi, Ratliff & de la Mora (2021), maps each of "
        "those quality outcomes to an alignment of instructional elements. GUIDE "
        "operationalizes each.",
    )
    elements_to_insert.append(p1)

    p2 = clone_para(
        existing_normal,
        "Effective - instructional elements aligned with theory, research, and "
        "documented best practice (Hannafin, Hannafin, Land, & Oliver, 1997). Whether "
        "the artifact reflects what learning science actually says works. This is the "
        "grounded in Grounded Universal Instructional Design Evaluator.",
    )
    elements_to_insert.append(p2)

    p3 = clone_para(
        existing_normal,
        "Efficient - objectives, instructional strategies (chunking and sequencing), "
        "and learner assessments aligned with each other (Tyler, 1949; Bloom, 1956; "
        "Dick, Carey & Carey, 2015). Whether the artifact's internal pieces cohere so "
        "learners are not doing wasted work. Archetype 10 (Curriculum Alignment) "
        "operationalizes this alignment as a dedicated evaluation lens.",
    )
    elements_to_insert.append(p3)

    p4 = clone_para(
        existing_normal,
        "Engaging - instructional elements aligned with learners' personal and "
        "professional goals, interests, and motivations (Keller, 1987, 2010 - ARCS). "
        "Whether the artifact connects to who the learner is and what they are trying "
        "to become. ARCS is operationalized as a capture requirement in Archetype 07 "
        "Level 4 (Learner Needs) and as a cross-cutting motivational design layer in "
        "Archetype 03 (Learning Domain Alignment).",
    )
    elements_to_insert.append(p4)

    p5 = clone_para(
        existing_normal,
        "Each archetype below evaluates one or more of these three alignments. "
        "Archetypes 01-09 evaluate component quality across the ADDIE lifecycle; "
        "Archetype 10 is the synthesis archetype that evaluates internal coherence "
        "between objectives, instruction, and assessment.",
    )
    elements_to_insert.append(p5)

    # Insert all elements before the Quick-Reference H2
    for el in elements_to_insert:
        insert_para_before(anchor, el)

    print(f"  inserted Quality Outcomes section ({len(elements_to_insert)} paragraphs)")


def update_quick_reference_table(doc):
    """Add a 10th row to the Archetype Quick-Reference table for Curriculum Alignment."""
    t = doc.tables[0]
    # Confirm it's the right table
    if t.rows[0].cells[0].text.strip() != "Archetype":
        print("  WARN: Table 0 is not the Quick-Reference table")
        return

    # Clone the last row to preserve formatting
    last_row = t.rows[-1]
    new_row_el = copy.deepcopy(last_row._element)
    last_row._element.addnext(new_row_el)

    # Set text in the new row's cells
    new_row = t.rows[-1]  # Now references the newly added row
    new_row.cells[0].text = "10 - Curriculum Alignment"
    new_row.cells[1].text = (
        "Objective <-> Strategy Coherence, Strategy <-> Assessment Coherence, "
        "Objective <-> Assessment Coherence, Coverage Completeness, Vertical "
        "Alignment, Discipline Alignment"
    )
    print("  Quick-Reference table: added Archetype 10 row")


def update_archetype_03_dim2(doc):
    """Augment Archetype 03 Dimension 2 (Learning Domain Alignment) with ARCS callout.

    Adds a new paragraph immediately after the existing explanation paragraph that
    introduces ARCS as a cross-cutting motivational layer.
    """
    # Find the H3 heading
    h3 = find_para(doc,
                   lambda style, text: style == "Heading 3"
                   and "Dimension 2: Learning Domain Alignment" in text)
    if h3 is None:
        print("  WARN: Archetype 03 D2 H3 not found")
        return

    # Walk forward to find the explanation paragraph (Normal, follows 'Theoretical Foundation' line)
    # then insert ARCS callout after it
    target = None
    el = h3._element
    seen_foundation = False
    while True:
        el = el.getnext()
        if el is None:
            break
        if el.tag.split("}")[-1] != "p":
            continue
        p = next((pp for pp in doc.paragraphs if pp._element is el), None)
        if p is None:
            continue
        if p.style.name.startswith("Heading"):
            break
        if p.text.startswith("Theoretical Foundation"):
            seen_foundation = True
            continue
        if seen_foundation and p.text.strip() and not p.text.startswith("Source Materials"):
            target = p
            break

    if target is None:
        print("  WARN: Archetype 03 D2 explanation paragraph not found")
        return

    # Insert ARCS callout
    new_p = clone_para(
        target,
        "ARCS Motivational Layer (added in v3.0.2). For the Attitudes domain - and as "
        "a cross-cutting motivational lens for any domain - apply Keller's ARCS: "
        "Attention (curiosity-grabbing entry), Relevance (tie to the learner's "
        "personal/professional goals captured in Archetype 07 Level 4), Confidence "
        "(early scaffolded wins, clear success criteria), Satisfaction (intrinsic "
        "reward, recognition, applied outcomes). If the upstream needs assessment "
        "captured ARCS data and this instruction ignores it in favor of generic "
        "framing, that is a domain-alignment failure on the motivational layer, not "
        "just a polish issue.",
    )
    insert_para_after(target, new_p)
    print("  Archetype 03 D2: ARCS callout inserted")


def update_archetype_07_dim3(doc):
    """Augment Archetype 07 Dimension 3 (Needs Assessment Completeness) Level 4
    requirements with ARCS (capture side)."""
    h3 = find_para(doc,
                   lambda style, text: style == "Heading 3"
                   and "Dimension 3: Needs Assessment Completeness" in text)
    if h3 is None:
        print("  WARN: Archetype 07 D3 H3 not found")
        return

    target = None
    el = h3._element
    seen_foundation = False
    while True:
        el = el.getnext()
        if el is None:
            break
        if el.tag.split("}")[-1] != "p":
            continue
        p = next((pp for pp in doc.paragraphs if pp._element is el), None)
        if p is None:
            continue
        if p.style.name.startswith("Heading"):
            break
        if p.text.startswith("Theoretical Foundation"):
            seen_foundation = True
            continue
        if seen_foundation and p.text.strip() and not p.text.startswith("Source Materials"):
            target = p
            break

    if target is None:
        print("  WARN: Archetype 07 D3 explanation paragraph not found")
        return

    new_p = clone_para(
        target,
        "ARCS at Level 4 (added in v3.0.2). Beyond demographics and constraints, "
        "Learner Needs must capture the four motivational conditions instruction must "
        "satisfy (Keller, 1987, 2010): Attention (what currently holds learners' "
        "interest? what competes for it?), Relevance (what personal or professional "
        "goals, interests, or pain points does this instruction connect to?), "
        "Confidence (current self-efficacy; where might they expect to fail?), and "
        "Satisfaction (what intrinsic or extrinsic outcomes will make the effort feel "
        "worth it?). A Level 4 profile that lists demographics but omits "
        "goals/interests/motivations scores no higher than 3 on this dimension - "
        "engagement is alignment to the learner's why, not just their context.",
    )
    insert_para_after(target, new_p)
    print("  Archetype 07 D3: ARCS L4 callout inserted")


def add_archetype_10(doc):
    """Insert the full Archetype 10: Curriculum Alignment section after Archetype 09
    and before the 'Severity Flagging' H1.

    Uses Archetype 09's structure as the template: H1 + description + 'When to use'
    1x1 table + 6 dimensions, each with H3 + Foundation + explanation + Source
    Materials + Scoring Criteria 6x2 table.
    """
    # Anchor: 'Severity Flagging' H1
    anchor = find_para(doc,
                       lambda style, text: style == "Heading 1" and "Severity Flagging" in text)
    if anchor is None:
        print("  WARN: 'Severity Flagging' anchor not found")
        return

    # Templates from Archetype 09
    a09_h1 = find_para(doc,
                       lambda style, text: style == "Heading 1"
                       and "Archetype 09" in text)
    if a09_h1 is None:
        print("  WARN: Archetype 09 H1 template not found")
        return

    # Find Archetype 09's 'When to use' 1x1 table (the first table after the A09 H1)
    when_to_use_table_template = None
    el = a09_h1._element
    while True:
        el = el.getnext()
        if el is None:
            break
        if el.tag.split("}")[-1] == "tbl":
            for t in doc.tables:
                if t._element is el:
                    when_to_use_table_template = t
                    break
            break
    if when_to_use_table_template is None:
        print("  WARN: A09 'When to use' table not found")
        return

    # Find a representative 6x2 scoring table (Archetype 09 Dimension 1 = table index 62)
    scoring_table_template = None
    for t in doc.tables:
        if len(t.rows) == 6 and len(t.rows[0].cells) == 2 and t.rows[0].cells[0].text.strip() == "Score":
            scoring_table_template = t
            break
    if scoring_table_template is None:
        print("  WARN: 6x2 scoring table template not found")
        return

    # Find an H3 paragraph (Dimension N), Normal paragraph, etc.
    h3_template = find_para(doc, lambda style, text: style == "Heading 3" and "Dimension 1" in text)
    normal_template = next((p for p in doc.paragraphs
                            if p.style.name == "Normal" and p.text.startswith("Theoretical Foundation")),
                           None)
    blank_template = next((p for p in doc.paragraphs
                           if p.style.name == "Normal" and not p.text.strip()),
                          None)
    if h3_template is None or normal_template is None:
        print("  WARN: H3 or Normal template not found")
        return

    # Build the Archetype 10 content
    insertions = []  # list of (kind, element) where kind in {'p', 't'}

    # H1
    insertions.append(("p", clone_para(a09_h1, "Archetype 10: Curriculum Alignment")))

    # Description paragraph (use a Normal template, this one is the A09 description)
    a09_desc = next((p for p in doc.paragraphs
                     if p.style.name == "Normal"
                     and p.text.startswith("Evaluates instruction for alignment")), None)
    desc_template = a09_desc if a09_desc is not None else normal_template
    insertions.append(("p", clone_para(
        desc_template,
        "Evaluates whether an instructional artifact's internal pieces - objectives, "
        "instructional strategies (content, activities, chunking, sequencing), and "
        "learner assessments - cohere with each other and with broader scopes "
        "(course-level vertical progression, discipline-level competencies). This "
        "archetype operationalizes the Efficient quality outcome (Merrill, 2002) per "
        "the alignment mapping in Hirumi (2025). It does not evaluate the quality of "
        "any individual component - other archetypes (02, 03, 07) do that. It "
        "evaluates whether the components, taken together, say the same thing.")))

    # 'When to use' 1x1 table - clone and set text
    when_to_use = clone_table(when_to_use_table_template)
    when_to_use.rows[0].cells[0].text = (
        "When to use this archetype:\n"
        "Apply for end-to-end course design reviews (objectives + lessons + "
        "assessments together), curriculum mapping against professional competencies, "
        "Bloom-drift checks between stated and tested behavior, detecting orphan "
        "objectives or extraneous instruction, and vertical alignment review across "
        "modules in a course or courses in a program. Run in addition to component "
        "archetypes (02, 03, 07) when the question is whether the pieces hang "
        "together rather than whether any one piece is well-designed."
    )
    insertions.append(("t", when_to_use))

    # Blank paragraph after table
    if blank_template is not None:
        insertions.append(("p", clone_para(blank_template, " ")))

    # Six dimensions
    dimensions = [
        {
            "title": "Dimension 1: Objective <-> Strategy Coherence",
            "foundation": "Theoretical Foundation: Tyler (1949) curriculum alignment; Dick, Carey & Carey (2015) objective-instruction linkage; Bloom (1956) / Anderson & Krathwohl (2001) taxonomy",
            "explanation": (
                "Does the instruction actually teach what the objective states? Behavior verb "
                "in the objective is practiced in instruction; Bloom level taught matches "
                "Bloom level stated; no instruction exists that does not trace to a stated "
                "objective."
            ),
            "sources": [
                "Tyler, R.W. (1949). Basic Principles of Curriculum and Instruction. University of Chicago Press.",
                "Dick, W., Carey, L., & Carey, J.O. (2015). The Systematic Design of Instruction (8th ed.). Pearson.",
                "Bloom, B.S. (1956). Taxonomy of Educational Objectives. David McKay.",
                "Anderson, L.W. & Krathwohl, D.R. (2001). A Taxonomy for Learning, Teaching, and Assessing. Longman.",
            ],
            "scores": [
                ("5", "Every objective's behavior verb is directly practiced in instruction. Bloom level taught matches Bloom level stated. No instruction exists that does not trace to an objective."),
                ("4", "Strong match; one minor verb-to-practice mismatch or one Bloom-level near-miss (one level off)."),
                ("3", "Moderate match; two or more dimensions teach a different cognitive level than the objective states (e.g., objective says 'analyze,' instruction stops at 'describe')."),
                ("2", "Significant misalignment; instruction teaches notably different skills than what the objectives state."),
                ("1", "Instruction and objectives are essentially disconnected; the artifact teaches one thing and claims another."),
            ],
        },
        {
            "title": "Dimension 2: Strategy <-> Assessment Coherence",
            "foundation": "Theoretical Foundation: Dick, Carey & Carey (2015) practice-assessment continuity; Mager (1997) congruence between practice and test conditions",
            "explanation": (
                "Does the assessment measure what was taught? Assessment tasks reflect the "
                "practice activities in format and cognitive demand; learners are not surprised "
                "by formats or content; the same Bloom level taught is tested."
            ),
            "sources": [
                "Dick, W., Carey, L., & Carey, J.O. (2015). The Systematic Design of Instruction (8th ed.). Pearson.",
                "Mager, R.F. (1997). Preparing Instructional Objectives (3rd ed.). Center for Effective Performance.",
            ],
            "scores": [
                ("5", "Assessment tasks mirror the practice activities in format and cognitive demand. No surprises for the learner."),
                ("4", "Mostly mirrored; one format change or one cognitive-demand step beyond what was practiced."),
                ("3", "Partial mirror; learners would face at least one major surprise in format or required reasoning."),
                ("2", "Significant mismatch; assessment requires skills that were not practiced."),
                ("1", "Assessment effectively tests different content than what was taught."),
            ],
        },
        {
            "title": "Dimension 3: Objective <-> Assessment Coherence",
            "foundation": "Theoretical Foundation: Mager (1997) verb-to-task matching; Bloom (1956) / Anderson & Krathwohl (2001) cognitive demand; Webb (1997) Depth of Knowledge",
            "explanation": (
                "Does the assessment test what the objective states? Verb-to-task match "
                "(Mager); Bloom-level match (Bloom / Anderson & Krathwohl); Webb's Depth of "
                "Knowledge consistent; the criterion stated in the objective is the criterion "
                "in the rubric. Bloom-level mismatch of 2 or more levels triggers a major "
                "severity flag."
            ),
            "sources": [
                "Mager, R.F. (1997). Preparing Instructional Objectives (3rd ed.). Center for Effective Performance.",
                "Bloom, B.S. (1956). Taxonomy of Educational Objectives. David McKay.",
                "Anderson, L.W. & Krathwohl, D.R. (2001). A Taxonomy for Learning, Teaching, and Assessing. Longman.",
                "Webb, N.L. (1997). Criteria for alignment of expectations and assessments. CCSSO Research Monograph No. 6.",
            ],
            "scores": [
                ("5", "Objective verb maps directly to assessment task (Mager); Bloom level of objective matches Bloom level of assessment item; Webb's DoK consistent; the criterion stated in the objective is the criterion in the rubric."),
                ("4", "Strong; one verb is slightly off (e.g., objective says 'evaluate,' assessment asks 'compare') OR one Bloom level near-miss."),
                ("3", "Two or more verb-to-task or Bloom-level mismatches."),
                ("2", "Multiple significant mismatches; assessment tests at a different cognitive level than stated objectives."),
                ("1", "Assessment tests fundamentally different behaviors than the objectives state. Overall is capped at 3.0 when this dimension scores 1 or 2."),
            ],
        },
        {
            "title": "Dimension 4: Coverage Completeness",
            "foundation": "Theoretical Foundation: Wiggins & McTighe (2005) Understanding by Design / Backward Design; Dick, Carey & Carey (2015) blueprint and table of specifications",
            "explanation": (
                "Is every objective served by instruction AND tested? Is every piece of "
                "instruction tied to an objective? A blueprint or table of specifications is "
                "present or trivially derivable. No orphan objectives (stated but not taught "
                "or not tested); no orphan instruction (taught but not aligned to any "
                "objective)."
            ),
            "sources": [
                "Wiggins, G., & McTighe, J. (2005). Understanding by Design (2nd ed.). ASCD.",
                "Dick, W., Carey, L., & Carey, J.O. (2015). The Systematic Design of Instruction (8th ed.). Pearson.",
            ],
            "scores": [
                ("5", "Every objective is addressed in both instruction and assessment; every piece of instruction traces to an objective. Blueprint or table of specifications present or trivially derivable. No orphans."),
                ("4", "Strong coverage; one minor orphan (one objective with no practice, OR one piece of instruction not tied to an objective)."),
                ("3", "Notable gaps; one stated objective is not assessed, OR substantial instruction is not tied to any objective."),
                ("2", "Multiple orphan objectives, or large blocks of extraneous instruction."),
                ("1", "Coverage is essentially incidental; objectives, instruction, and assessment are three independent lists."),
            ],
        },
        {
            "title": "Dimension 5: Vertical Alignment",
            "foundation": "Theoretical Foundation: Bruner spiral curriculum; Reigeluth (1999) Elaboration Theory; van Merrienboer & Kirschner (2018) 4C/ID whole-task progression",
            "explanation": (
                "Across lessons, modules, or courses, do they build coherently? Prerequisites "
                "identified and respected. Spiral curriculum (Bruner) or elaboration logic "
                "(Reigeluth) evident. Later content extends, not merely repeats, earlier "
                "content. For single-lesson artifacts, score within-lesson sequencing instead "
                "and flag the adaptation in rationale."
            ),
            "sources": [
                "Reigeluth, C.M. (1999). Instructional-Design Theories and Models, Volume II. Lawrence Erlbaum.",
                "van Merrienboer, J.J.G. & Kirschner, P.A. (2018). Ten Steps to Complex Learning (3rd ed.). Routledge.",
            ],
            "scores": [
                ("5", "Across lessons/modules/courses, content builds coherently. Prerequisites explicit and respected. Spiral or elaboration logic evident. Later content extends rather than repeats."),
                ("4", "Strong vertical logic; one prerequisite gap or one redundancy."),
                ("3", "Moderate; some modules build on prior, others stand alone in ways that do not make pedagogical sense."),
                ("2", "Weak vertical logic; modules feel like independent units rather than a connected progression."),
                ("1", "No vertical alignment; lesson/module order is arbitrary or reverse-coherent."),
            ],
        },
        {
            "title": "Dimension 6: Discipline Alignment",
            "foundation": "Theoretical Foundation: Hirumi, Ratliff & de la Mora (2021) discipline-scope alignment in Figure 1; Hirumi (2025) three-alignments mapping; competency-based education traditions (IBSTPI, ATD, Quality Matters)",
            "explanation": (
                "Does the course or program ladder up to professional competencies, knowledge "
                "base, or accrediting examinations? Every objective at the unit/lesson level "
                "maps upward to a competency at the discipline level. For artifacts not "
                "intended to ladder up to a discipline (e.g., one-off workshops, internal "
                "training), score on internal coherence with stated purpose and note the "
                "adaptation in rationale."
            ),
            "sources": [
                "Hirumi, A., Ratliff, M., & de la Mora Velasco, E. (2021). Analyzing your context to improve and innovate distance learning. In L. Cifuentes (Ed.), Guide to Administering Distance Learning. Brill.",
                "Hirumi, A. (2025). Design Principles, Ethics, and Evidence-Informed Decision Making. In R.E. West & H. Leary (Eds.), Foundations of Learning and Instructional Design Technology (2nd ed., pp. 433-448). EdTechBooks.",
            ],
            "scores": [
                ("5", "Explicit, traceable linkage to external competency framework (national exam, professional standards, accreditation body). Every objective at the unit/lesson level maps upward to a competency at the discipline level."),
                ("4", "Strong linkage; one objective does not map cleanly to a stated competency."),
                ("3", "Linkage stated but not traceable in detail; competency framework named but specific mappings unclear."),
                ("2", "Weak linkage; some content seems to serve external standards but the mapping is left to the reader."),
                ("1", "No stated discipline-level alignment; course exists in isolation from professional competencies or accreditation."),
            ],
        },
    ]

    for dim in dimensions:
        insertions.append(("p", clone_para(h3_template, dim["title"])))
        insertions.append(("p", clone_para(normal_template, dim["foundation"])))
        insertions.append(("p", clone_para(normal_template, dim["explanation"])))
        insertions.append(("p", clone_para(normal_template, "Source Materials:")))
        for src in dim["sources"]:
            insertions.append(("p", clone_para(normal_template, src)))
        if blank_template is not None:
            insertions.append(("p", clone_para(blank_template, " ")))
        insertions.append(("p", clone_para(normal_template, "Scoring Criteria:")))

        # Build a fresh scoring table from the template
        st = clone_table(scoring_table_template)
        # Header row stays "Score | Criteria"; populate rows 1-5
        for i, (score, criterion) in enumerate(dim["scores"], start=1):
            st.rows[i].cells[0].text = score
            st.rows[i].cells[1].text = criterion
        insertions.append(("t", st))
        if blank_template is not None:
            insertions.append(("p", clone_para(blank_template, " ")))

    # Insert everything before the 'Severity Flagging' anchor.
    # We insert each element in order using addprevious() on the anchor.
    for kind, elem_wrapper in insertions:
        anchor._element.addprevious(elem_wrapper._element)

    print(f"  Archetype 10: inserted {len(insertions)} body elements (paragraphs + tables)")


def update_references(doc):
    """Append the new citations to the References section."""
    refs_h1 = find_para(doc,
                        lambda style, text: style == "Heading 1" and text.strip() == "References")
    if refs_h1 is None:
        print("  WARN: References section H1 not found")
        return

    # Walk forward and gather existing Normal paragraphs to find a template
    normal_template = None
    last_ref_para = None
    el = refs_h1._element
    while True:
        el = el.getnext()
        if el is None:
            break
        if el.tag.split("}")[-1] != "p":
            continue
        p = next((pp for pp in doc.paragraphs if pp._element is el), None)
        if p is None:
            continue
        if p.style.name.startswith("Heading"):
            break
        if p.style.name == "Normal" and p.text.strip():
            normal_template = p
            last_ref_para = p

    if normal_template is None:
        print("  WARN: no Normal reference paragraph template found")
        return

    new_refs = [
        "",
        "References added in v3.1.0 - v3.1.2:",
        "Tyler, R.W. (1949). Basic Principles of Curriculum and Instruction. University of Chicago Press.",
        "Anderson, L.W. & Krathwohl, D.R. (Eds.) (2001). A Taxonomy for Learning, Teaching, and Assessing: A Revision of Bloom's Taxonomy of Educational Objectives. Longman.",
        "Webb, N.L. (1997). Criteria for alignment of expectations and assessments in mathematics and science education. CCSSO Research Monograph No. 6.",
        "Wiggins, G., & McTighe, J. (2005). Understanding by Design (2nd ed.). ASCD.",
        "Reigeluth, C.M. (1999). Instructional-Design Theories and Models, Volume II: A New Paradigm of Instructional Theory. Lawrence Erlbaum.",
        "van Merrienboer, J.J.G., & Kirschner, P.A. (2018). Ten Steps to Complex Learning (3rd ed.). Routledge.",
        "Merrill, M.D. (2002). First principles of instruction. Educational Technology Research and Development, 50(3), 43-59.",
        "Keller, J.M. (1987). Development and Use of the ARCS Model of Instructional Design. Journal of Instructional Development, 10(3), 2-10.",
        "Keller, J.M. (2010). Motivational Design for Learning and Performance: The ARCS Model Approach. Springer.",
        "Hannafin, M.J., Hannafin, K.M., Land, S.M., & Oliver, K. (1997). Grounded practice and the design of constructivist learning environments. Educational Technology Research and Development, 45(3), 101-117.",
        "Hirumi, A., Ratliff, M., & de la Mora Velasco, E. (2021). Analyzing your context to improve and innovate distance learning (pp. 104-140). In L. Cifuentes (Ed.), Guide to Administering Distance Learning. Brill Publishing.",
        "Hirumi, A. (2025). Design Principles, Ethics, and Evidence-Informed Decision Making: The (Evolving) Future of the Field is Design! In R.E. West & H. Leary (Eds.), Foundations of Learning and Instructional Design Technology (2nd ed., pp. 244-254 and 433-448). EdTechBooks.",
    ]

    # Insert each as a clone of the template, after the last reference paragraph
    target = last_ref_para
    for ref_text in new_refs:
        new_p = clone_para(normal_template, ref_text if ref_text else " ")
        insert_para_after(target, new_p)
        target = new_p
    print(f"  References: appended {len(new_refs)} entries")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    if not DOCX_PATH.exists():
        raise SystemExit(f"Cannot find {DOCX_PATH}")

    doc = Document(str(DOCX_PATH))
    print(f"Loaded {DOCX_PATH} ({len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables)")

    update_cover_version(doc)
    update_executive_summary(doc)
    update_why_n_archetypes(doc)
    add_quality_outcomes_section(doc)
    update_quick_reference_table(doc)
    update_archetype_03_dim2(doc)
    update_archetype_07_dim3(doc)
    add_archetype_10(doc)
    update_references(doc)

    doc.save(str(DOCX_PATH))
    print(f"Saved {DOCX_PATH}")


if __name__ == "__main__":
    main()
